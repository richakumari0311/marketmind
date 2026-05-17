import asyncio
import json
import os
import re
from datetime import datetime
from io import BytesIO

from fastapi import FastAPI
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse, JSONResponse, StreamingResponse
from sse_starlette.sse import EventSourceResponse
from pydantic import BaseModel

from agents.research_agent import run as research
from agents.strategy_agent import run as strategy
from agents.content_agent import run as write_content
from agents.critique_agent import run as critique
from utils import memory as mem
from utils.notion_writer import write_campaign

app = FastAPI()
app.mount("/static", StaticFiles(directory="static"), name="static")


class CampaignRequest(BaseModel):
    product: str
    goal: str = "awareness"
    pieces: int = 3


@app.get("/")
def index():
    return FileResponse("static/index.html")


@app.get("/campaigns")
def get_campaigns():
    return JSONResponse(mem.list_campaigns())


@app.get("/campaigns/{run_id}")
def get_campaign(run_id: str):
    data = mem.load_campaign_file(run_id)
    if not data:
        return JSONResponse({"error": "Campaign not found"}, status_code=404)
    return JSONResponse(data)


@app.get("/campaigns/{run_id}/export/pdf")
def export_pdf(run_id: str):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.lib import colors
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer,
        Table, TableStyle, HRFlowable
    )

    data = mem.load_campaign_file(run_id)
    if not data:
        return JSONResponse({"error": "Campaign not found"}, status_code=404)

    strategy_data = data.get("strategy", {})
    critique_data = data.get("critique", {})
    content_data  = data.get("content", {})
    meta          = data.get("meta", {})

    buf = BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=20*mm, rightMargin=20*mm,
        topMargin=20*mm, bottomMargin=20*mm
    )

    styles = getSampleStyleSheet()
    brand  = colors.HexColor("#2563eb")

    title_style = ParagraphStyle(
        "Title", parent=styles["Title"],
        textColor=brand, fontSize=24, spaceAfter=4
    )
    h1_style = ParagraphStyle(
        "H1", parent=styles["Heading1"],
        textColor=brand, fontSize=14, spaceBefore=16, spaceAfter=6
    )
    h2_style = ParagraphStyle(
        "H2", parent=styles["Heading2"],
        fontSize=11, spaceBefore=10, spaceAfter=4
    )
    body_style = ParagraphStyle(
        "Body", parent=styles["Normal"],
        fontSize=9, leading=14, spaceAfter=6
    )
    meta_style = ParagraphStyle(
        "Meta", parent=styles["Normal"],
        fontSize=9, textColor=colors.HexColor("#64748b")
    )

    def clean(text):
        if not text:
            return ""
        text = str(text)
        text = text.replace("\\n", "\n")
        text = re.sub(r"\\u([0-9a-fA-F]{4})", lambda m: chr(int(m.group(1), 16)), text)
        text = re.sub(r"<[^>]+>", "", text)
        return text.strip()

    story = []

    story.append(Paragraph(clean(strategy_data.get("campaign_name", "Campaign")), title_style))
    story.append(Paragraph(
        f"Goal: {meta.get('goal','').upper()}  |  "
        f"Date: {meta.get('date','')}  |  "
        f"Avg Score: {meta.get('avg_score',0)}/100",
        meta_style
    ))
    story.append(Spacer(1, 6))
    story.append(HRFlowable(width="100%", color=brand))
    story.append(Spacer(1, 12))

    story.append(Paragraph("Core Message", h1_style))
    story.append(Paragraph(clean(strategy_data.get("core_message", "")), body_style))

    mh = strategy_data.get("messaging_hierarchy", {})
    if mh:
        story.append(Paragraph("Messaging Hierarchy", h1_style))
        story.append(Paragraph(f"<b>Headline:</b> {clean(mh.get('headline',''))}", body_style))
        story.append(Paragraph(f"<b>Subheadline:</b> {clean(mh.get('subheadline',''))}", body_style))
        story.append(Paragraph(f"<b>CTA:</b> {clean(mh.get('cta',''))}", body_style))
        for p in mh.get("proof_points", []):
            story.append(Paragraph(f"- {clean(p)}", body_style))

    kpis = strategy_data.get("kpis", [])
    if kpis:
        story.append(Paragraph("KPIs", h1_style))
        table_data = [["Metric", "Target", "Timeframe"]]
        for k in kpis:
            table_data.append([
                clean(k.get("metric", "")),
                clean(k.get("target", "")),
                clean(k.get("timeframe", "")),
            ])
        t = Table(table_data, colWidths=[65*mm, 75*mm, 35*mm])
        t.setStyle(TableStyle([
            ("BACKGROUND",     (0,0), (-1,0), brand),
            ("TEXTCOLOR",      (0,0), (-1,0), colors.white),
            ("FONTSIZE",       (0,0), (-1,-1), 9),
            ("ROWBACKGROUNDS", (0,1), (-1,-1), [colors.white, colors.HexColor("#f1f5f9")]),
            ("GRID",           (0,0), (-1,-1), 0.25, colors.HexColor("#e2e8f0")),
            ("TOPPADDING",     (0,0), (-1,-1), 4),
            ("BOTTOMPADDING",  (0,0), (-1,-1), 4),
        ]))
        story.append(t)

    scores = critique_data.get("scores", [])
    if scores:
        story.append(Paragraph("Content Scores", h1_style))
        score_data = [["Title", "Clarity", "Relevance", "Conversion", "Originality", "Total"]]
        for s in scores:
            dims = s.get("scores", {})
            score_data.append([
                clean(s.get("title", ""))[:40],
                str(dims.get("clarity", 0)),
                str(dims.get("relevance", 0)),
                str(dims.get("conversion", 0)),
                str(dims.get("originality", 0)),
                str(s.get("total", 0)),
            ])
        t = Table(score_data, colWidths=[70*mm, 20*mm, 22*mm, 25*mm, 25*mm, 15*mm])
        t.setStyle(TableStyle([
            ("BACKGROUND",     (0,0), (-1,0), brand),
            ("TEXTCOLOR",      (0,0), (-1,0), colors.white),
            ("FONTSIZE",       (0,0), (-1,-1), 9),
            ("ROWBACKGROUNDS", (0,1), (-1,-1), [colors.white, colors.HexColor("#f1f5f9")]),
            ("GRID",           (0,0), (-1,-1), 0.25, colors.HexColor("#e2e8f0")),
            ("TOPPADDING",     (0,0), (-1,-1), 4),
            ("BOTTOMPADDING",  (0,0), (-1,-1), 4),
        ]))
        story.append(t)

    pieces = content_data.get("content", [])
    if pieces:
        story.append(Paragraph("Content Pieces", h1_style))
        for piece in pieces:
            story.append(Paragraph(clean(piece.get("title", "Untitled")), h2_style))
            story.append(Paragraph(
                f"{piece.get('type','').upper()}  |  "
                f"{piece.get('channel','')}  |  "
                f"{piece.get('meta',{}).get('word_count',0)} words",
                meta_style
            ))
            story.append(Spacer(1, 4))
            story.append(Paragraph(clean(piece.get("content", ""))[:2000], body_style))
            cta = piece.get("meta", {}).get("cta", "")
            if cta:
                story.append(Paragraph(f"<b>CTA:</b> {clean(cta)}", body_style))
            story.append(Spacer(1, 8))

    rewritten = critique_data.get("weakest_rewritten")
    if rewritten:
        story.append(HRFlowable(width="100%", color=colors.HexColor("#16a34a")))
        story.append(Paragraph("Auto-Rewritten Piece", h1_style))
        story.append(Paragraph(clean(rewritten.get("title", "")), h2_style))
        story.append(Paragraph(clean(rewritten.get("content", ""))[:2000], body_style))
        for c in rewritten.get("changes", []):
            story.append(Paragraph(f"+ {clean(c)}", body_style))

    doc.build(story)
    buf.seek(0)

    filename = f"marketmind_{run_id}.pdf"
    return StreamingResponse(
        buf,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )


@app.get("/campaigns/{run_id}/export/docx")
def export_docx(run_id: str):
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    data = mem.load_campaign_file(run_id)
    if not data:
        return JSONResponse({"error": "Campaign not found"}, status_code=404)

    strategy_data = data.get("strategy", {})
    critique_data = data.get("critique", {})
    content_data  = data.get("content", {})
    meta          = data.get("meta", {})

    def clean(text):
        if not text:
            return ""
        text = str(text)
        text = text.replace("\\n", "\n")
        text = re.sub(r"\\u([0-9a-fA-F]{4})", lambda m: chr(int(m.group(1), 16)), text)
        text = re.sub(r"<[^>]+>", "", text)
        return text.strip()

    BRAND = RGBColor(0x25, 0x63, 0xEB)
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    def add_heading(text, level=1, color=None):
        p = doc.add_heading(clean(text), level=level)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if color:
            for run in p.runs:
                run.font.color.rgb = color
        return p

    def add_body(text):
        p = doc.add_paragraph(clean(text))
        p.style.font.size = Pt(10)
        return p

    def add_meta(text):
        p = doc.add_paragraph(clean(text))
        p.style.font.size = Pt(9)
        return p

    title = doc.add_heading(clean(strategy_data.get("campaign_name", "Campaign")), 0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in title.runs:
        run.font.color.rgb = BRAND

    add_meta(
        f"Goal: {meta.get('goal','').upper()}  |  "
        f"Date: {meta.get('date','')}  |  "
        f"Avg Score: {meta.get('avg_score',0)}/100"
    )
    doc.add_paragraph()

    add_heading("Core Message", 1, BRAND)
    add_body(strategy_data.get("core_message", ""))

    mh = strategy_data.get("messaging_hierarchy", {})
    if mh:
        add_heading("Messaging Hierarchy", 1, BRAND)
        add_body(f"Headline: {mh.get('headline','')}")
        add_body(f"Subheadline: {mh.get('subheadline','')}")
        add_body(f"CTA: {mh.get('cta','')}")
        for point in mh.get("proof_points", []):
            add_body(f"- {point}")

    kpis = strategy_data.get("kpis", [])
    if kpis:
        add_heading("KPIs", 1, BRAND)
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Metric"
        hdr[1].text = "Target"
        hdr[2].text = "Timeframe"
        for k in kpis:
            row = table.add_row().cells
            row[0].text = clean(k.get("metric", ""))
            row[1].text = clean(k.get("target", ""))
            row[2].text = clean(k.get("timeframe", ""))

    scores = critique_data.get("scores", [])
    if scores:
        add_heading("Content Scores", 1, BRAND)
        table = doc.add_table(rows=1, cols=6)
        table.style = "Table Grid"
        headers = ["Title", "Clarity", "Relevance", "Conversion", "Originality", "Total"]
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
        for s in scores:
            dims = s.get("scores", {})
            row = table.add_row().cells
            row[0].text = clean(s.get("title", ""))[:40]
            row[1].text = str(dims.get("clarity", 0))
            row[2].text = str(dims.get("relevance", 0))
            row[3].text = str(dims.get("conversion", 0))
            row[4].text = str(dims.get("originality", 0))
            row[5].text = str(s.get("total", 0))

    pieces = content_data.get("content", [])
    if pieces:
        add_heading("Content Pieces", 1, BRAND)
        for piece in pieces:
            add_heading(piece.get("title", "Untitled"), 2)
            add_meta(
                f"{piece.get('type','').upper()}  |  "
                f"{piece.get('channel','')}  |  "
                f"{piece.get('meta',{}).get('word_count',0)} words"
            )
            add_body(piece.get("content", "")[:2000])
            cta = piece.get("meta", {}).get("cta", "")
            if cta:
                add_body(f"CTA: {cta}")
            doc.add_paragraph()

    rewritten = critique_data.get("weakest_rewritten")
    if rewritten:
        add_heading("Auto-Rewritten Piece", 1, BRAND)
        add_heading(rewritten.get("title", ""), 2)
        add_body(rewritten.get("content", "")[:2000])
        for c in rewritten.get("changes", []):
            add_body(f"+ {clean(c)}")

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    filename = f"marketmind_{run_id}.docx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )


@app.post("/run")
async def run_campaign(req: CampaignRequest):

    async def stream():
        loop = asyncio.get_event_loop()

        def emit(event_type: str, data: dict):
            return {"event": event_type, "data": json.dumps(data)}

        research_output = None
        strategy_output = None
        content_output  = None
        run_id = datetime.now().strftime("%Y%m%d_%H%M%S")

        try:
            yield emit("progress", {"phase": 1, "label": "ResearchAgent", "status": "running"})
            research_output = await loop.run_in_executor(None, research, req.product)
            yield emit("progress", {"phase": 1, "label": "ResearchAgent", "status": "done"})
            yield emit("research", research_output)
        except Exception as e:
            yield emit("progress", {"phase": 1, "label": "ResearchAgent", "status": "error"})
            yield emit("error", {"phase": 1, "message": f"ResearchAgent failed: {str(e)}"})
            return

        try:
            yield emit("progress", {"phase": 2, "label": "StrategyAgent", "status": "running"})
            strategy_output = await loop.run_in_executor(
                None, lambda: strategy(research_output, campaign_goal=req.goal)
            )
            yield emit("progress", {"phase": 2, "label": "StrategyAgent", "status": "done"})
            yield emit("strategy", strategy_output)
        except Exception as e:
            yield emit("progress", {"phase": 2, "label": "StrategyAgent", "status": "error"})
            yield emit("error", {"phase": 2, "message": f"StrategyAgent failed: {str(e)}"})
            return

        try:
            yield emit("progress", {"phase": 3, "label": "ContentAgent", "status": "running"})
            content_output = await loop.run_in_executor(
                None, lambda: write_content(strategy_output, max_pieces=req.pieces)
            )
            yield emit("progress", {"phase": 3, "label": "ContentAgent", "status": "done"})
            yield emit("content", content_output)
        except Exception as e:
            yield emit("progress", {"phase": 3, "label": "ContentAgent", "status": "error"})
            yield emit("error", {"phase": 3, "message": f"ContentAgent failed: {str(e)}"})
            yield emit("done", {
                "campaign_name": strategy_output.get("campaign_name", ""),
                "partial": True
            })
            return

        try:
            yield emit("progress", {"phase": 4, "label": "CritiqueAgent", "status": "running"})
            critique_output = await loop.run_in_executor(
                None, lambda: critique(content_output, strategy_output)
            )
            yield emit("progress", {"phase": 4, "label": "CritiqueAgent", "status": "done"})
            yield emit("critique", critique_output)
        except Exception as e:
            yield emit("progress", {"phase": 4, "label": "CritiqueAgent", "status": "error"})
            yield emit("error", {"phase": 4, "message": f"CritiqueAgent failed: {str(e)}"})
            yield emit("done", {
                "campaign_name": strategy_output.get("campaign_name", ""),
                "partial": True
            })
            return

        avg_score = round(
            sum(s.get("total", 0) for s in critique_output.get("scores", [])) /
            max(len(critique_output.get("scores", [])), 1), 1
        )

        full_data = {
            "research": research_output,
            "strategy": strategy_output,
            "content":  content_output,
            "critique": critique_output,
            "meta": {
                "run_id":           run_id,
                "product":          req.product,
                "goal":             req.goal,
                "date":             datetime.now().strftime("%Y-%m-%d %H:%M"),
                "campaign_name":    strategy_output.get("campaign_name", ""),
                "pieces_generated": req.pieces,
                "avg_score":        avg_score,
                "total_time_seconds": 0,
            }
        }

        mem.record_campaign(
            product=req.product,
            campaign_name=strategy_output.get("campaign_name", ""),
            research=research_output,
            strategy=strategy_output,
            scores=critique_output.get("scores", []),
        )
        mem.save_campaign_file(run_id, full_data)

        notion_url = ""
        try:
            notion_url = await loop.run_in_executor(
                None, lambda: write_campaign(full_data)
            )
        except Exception as e:
            print(f"  [WARN] Notion write failed: {e}")

        yield emit("done", {
            "campaign_name": strategy_output.get("campaign_name", ""),
            "run_id":        run_id,
            "avg_score":     avg_score,
            "notion_url":    notion_url,
            "partial":       False,
        })

    return EventSourceResponse(stream())